#!/usr/bin/env python3
"""
Logging & Reporting Module for PayBuddy Cybersecurity Testing Toolkit
Append-only logs with SHA-256 integrity and auto-generate Word/PDF reports
"""

import json
import hashlib
import os
import sys
import time
from datetime import datetime
from pathlib import Path
import logging
import argparse

# Add src directory to path for imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from identity_verifier import IdentityVerifier

# Optional imports for report generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class SecurityLogger:
    def __init__(self, log_directory=None):
        if log_directory is None:
            # Default to evidence directory
            self.log_directory = Path(__file__).parent.parent / "evidence"
        else:
            self.log_directory = Path(log_directory)
        
        self.log_directory.mkdir(exist_ok=True)
        
        # Main log file
        self.log_file = self.log_directory / "security_audit.log"
        self.integrity_file = self.log_directory / "integrity.sha256"
        
        # Setup logging
        self.setup_logging()
        
        # Initialize integrity tracking
        self.init_integrity()
    
    def setup_logging(self):
        """Setup logging configuration"""
        # Create custom formatter
        formatter = logging.Formatter(
            '%(asctime)s | %(levelname)s | %(module)s | %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        
        # File handler for append-only logging
        file_handler = logging.FileHandler(self.log_file, mode='a')
        file_handler.setFormatter(formatter)
        file_handler.setLevel(logging.INFO)
        
        # Create logger
        self.logger = logging.getLogger('paybuddy_security')
        self.logger.setLevel(logging.INFO)
        
        # Clear any existing handlers and add our custom handler
        self.logger.handlers = []
        self.logger.addHandler(file_handler)
        
        # Prevent propagation to root logger
        self.logger.propagate = False
    
    def init_integrity(self):
        """Initialize integrity checking"""
        if self.log_file.exists():
            # Calculate current hash
            current_hash = self.calculate_file_hash(self.log_file)
            
            # Check if integrity file exists
            if self.integrity_file.exists():
                try:
                    with open(self.integrity_file, 'r') as f:
                        stored_hash = f.read().strip()
                    
                    if stored_hash != current_hash:
                        print("⚠️  Warning: Log file integrity check failed!")
                        print("   This may indicate tampering or corruption.")
                except:
                    pass
            
            # Update integrity file
            self.update_integrity()
        else:
            # Create empty log file and initial integrity
            self.log_file.touch()
            self.update_integrity()
    
    def calculate_file_hash(self, file_path):
        """Calculate SHA-256 hash of a file"""
        sha256_hash = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                for byte_block in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(byte_block)
        except:
            return ""
        return sha256_hash.hexdigest()
    
    def update_integrity(self):
        """Update integrity hash"""
        if self.log_file.exists():
            current_hash = self.calculate_file_hash(self.log_file)
            with open(self.integrity_file, 'w') as f:
                f.write(current_hash)
    
    def log_event(self, event_type, message, module="general", data=None):
        """Log a security event"""
        # Prepare log entry
        log_data = {
            'timestamp': datetime.now().isoformat(),
            'event_type': event_type,
            'module': module,
            'message': message,
            'data': data
        }
        
        # Create formatted message
        formatted_message = f"{event_type} | {message}"
        if data:
            formatted_message += f" | Data: {json.dumps(data, separators=(',', ':'))}"
        
        # Log to file
        self.logger.info(formatted_message)
        
        # Update integrity hash after each write
        self.update_integrity()
        
        return log_data
    
    def log_test_start(self, test_name, target, config=None):
        """Log the start of a security test"""
        return self.log_event(
            'TEST_START',
            f"Starting {test_name} against {target}",
            module=test_name.lower().replace(' ', '_'),
            data={
                'target': target,
                'config': config,
                'start_time': time.time()
            }
        )
    
    def log_test_end(self, test_name, target, results=None, duration=None):
        """Log the end of a security test"""
        return self.log_event(
            'TEST_END',
            f"Completed {test_name} against {target}",
            module=test_name.lower().replace(' ', '_'),
            data={
                'target': target,
                'results': results,
                'duration': duration,
                'end_time': time.time()
            }
        )
    
    def log_finding(self, severity, finding_type, description, target=None, evidence=None):
        """Log a security finding"""
        return self.log_event(
            'FINDING',
            f"[{severity.upper()}] {finding_type}: {description}",
            module='findings',
            data={
                'severity': severity,
                'type': finding_type,
                'target': target,
                'evidence': evidence
            }
        )
    
    def log_error(self, error_message, module="unknown", details=None):
        """Log an error"""
        return self.log_event(
            'ERROR',
            error_message,
            module=module,
            data=details
        )
    
    def get_log_entries(self, filter_by=None, start_time=None, end_time=None):
        """Retrieve log entries with optional filtering"""
        entries = []
        
        try:
            with open(self.log_file, 'r') as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    
                    try:
                        # Parse log line
                        parts = line.split(' | ')
                        if len(parts) >= 4:
                            entry = {
                                'timestamp': parts[0],
                                'level': parts[1],
                                'module': parts[2],
                                'message': ' | '.join(parts[3:])
                            }
                            
                            # Apply filters
                            if filter_by and filter_by not in entry['message']:
                                continue
                            
                            # TODO: Add time filtering if needed
                            
                            entries.append(entry)
                    except:
                        continue
        except FileNotFoundError:
            pass
        
        return entries

class ReportGenerator:
    def __init__(self, logger):
        self.logger = logger
        self.verifier = IdentityVerifier()
    
    def generate_summary_report(self, output_file, format='json'):
        """Generate a summary report of all security testing activities"""
        # Get team information
        team_info = self.verifier.get_team_info()
        
        # Collect log entries
        log_entries = self.logger.get_log_entries()
        
        # Analyze logs
        analysis = self.analyze_logs(log_entries)
        
        # Generate report based on format
        if format.lower() == 'json':
            return self.generate_json_report(output_file, team_info, analysis, log_entries)
        elif format.lower() == 'txt':
            return self.generate_txt_report(output_file, team_info, analysis, log_entries)
        elif format.lower() == 'docx' and DOCX_AVAILABLE:
            return self.generate_word_report(output_file, team_info, analysis, log_entries)
        elif format.lower() == 'pdf' and PDF_AVAILABLE:
            return self.generate_pdf_report(output_file, team_info, analysis, log_entries)
        else:
            print(f"❌ Format '{format}' not supported or missing dependencies")
            return False
    
    def analyze_logs(self, log_entries):
        """Analyze log entries to generate statistics"""
        analysis = {
            'total_entries': len(log_entries),
            'test_counts': {},
            'findings': [],
            'errors': [],
            'timeline': [],
            'modules_used': set(),
            'test_summary': {
                'total_tests': 0,
                'successful_tests': 0,
                'failed_tests': 0
            }
        }
        
        for entry in log_entries:
            # Extract module
            analysis['modules_used'].add(entry['module'])
            
            # Analyze message content
            message = entry['message']
            
            if 'TEST_START' in message:
                analysis['test_summary']['total_tests'] += 1
                test_name = message.split('Starting ')[1].split(' against')[0] if 'Starting ' in message else 'Unknown'
                analysis['test_counts'][test_name] = analysis['test_counts'].get(test_name, 0) + 1
            
            elif 'TEST_END' in message:
                analysis['test_summary']['successful_tests'] += 1
            
            elif 'FINDING' in message:
                # Extract finding information
                if '[HIGH]' in message or '[CRITICAL]' in message:
                    finding = {
                        'severity': 'HIGH' if '[HIGH]' in message else 'CRITICAL',
                        'message': message,
                        'timestamp': entry['timestamp']
                    }
                    analysis['findings'].append(finding)
            
            elif 'ERROR' in message:
                error = {
                    'message': message,
                    'timestamp': entry['timestamp'],
                    'module': entry['module']
                }
                analysis['errors'].append(error)
            
            # Build timeline
            analysis['timeline'].append({
                'timestamp': entry['timestamp'],
                'type': 'TEST_START' if 'TEST_START' in message else 
                       'TEST_END' if 'TEST_END' in message else
                       'FINDING' if 'FINDING' in message else
                       'ERROR' if 'ERROR' in message else 'INFO',
                'message': message[:100] + ('...' if len(message) > 100 else '')
            })
        
        # Convert set to list for JSON serialization
        analysis['modules_used'] = list(analysis['modules_used'])
        
        return analysis
    
    def generate_json_report(self, output_file, team_info, analysis, log_entries):
        """Generate JSON format report"""
        report = {
            'report_info': {
                'generated_at': datetime.now().isoformat(),
                'report_version': 'PayBuddy Security Report v1.0',
                'team_info': team_info,
                'log_file_integrity': self.logger.calculate_file_hash(self.logger.log_file)
            },
            'analysis': analysis,
            'raw_logs': log_entries[-100:]  # Last 100 entries to keep size manageable
        }
        
        try:
            # Ensure the directory exists
            from pathlib import Path
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_file, 'w') as f:
                json.dump(report, f, indent=2)
            print(f"📄 JSON report generated: {output_file}")
            return True
        except Exception as e:
            print(f"❌ Error generating JSON report: {e}")
            return False
    
    def generate_txt_report(self, output_file, team_info, analysis, log_entries):
        """Generate plain text format report"""
        try:
            # Ensure the directory exists
            from pathlib import Path
            output_path = Path(output_file)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(output_file, 'w', encoding='utf-8') as f:
                # Header
                f.write("=" * 80 + "\n")
                f.write("PayBuddy Cybersecurity Assessment Report\n")
                f.write("=" * 80 + "\n\n")
                
                # Report info
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Report Version: PayBuddy Security Report v1.0\n\n")
                
                # Team Information
                f.write("TEAM INFORMATION\n")
                f.write("-" * 40 + "\n")
                f.write(f"Team: {team_info['team']}\n")
                f.write("Members:\n")
                for member in team_info['members']:
                    f.write(f"  - {member}\n")
                f.write("\n")
                
                # Analysis Summary
                f.write("ANALYSIS SUMMARY\n")
                f.write("-" * 40 + "\n")
                f.write(f"Total Log Entries: {analysis['total_entries']}\n")
                f.write(f"Modules Used: {', '.join(analysis['modules_used'])}\n")
                f.write(f"Tests Performed: {len(analysis['test_counts'])}\n")
                f.write(f"Findings: {len(analysis['findings'])}\n")
                f.write(f"Errors: {len(analysis['errors'])}\n\n")
                
                # Test Counts
                if analysis['test_counts']:
                    f.write("TEST COUNTS\n")
                    f.write("-" * 40 + "\n")
                    for test, count in analysis['test_counts'].items():
                        f.write(f"  {test}: {count}\n")
                    f.write("\n")
                
                # Findings
                if analysis['findings']:
                    f.write("SECURITY FINDINGS\n")
                    f.write("-" * 40 + "\n")
                    for i, finding in enumerate(analysis['findings'], 1):
                        f.write(f"{i}. {finding}\n")
                    f.write("\n")
                
                # Errors
                if analysis['errors']:
                    f.write("ERRORS ENCOUNTERED\n")
                    f.write("-" * 40 + "\n")
                    for i, error in enumerate(analysis['errors'], 1):
                        f.write(f"{i}. {error}\n")
                    f.write("\n")
                
                # Recent Log Entries (last 20)
                f.write("RECENT ACTIVITY LOG\n")
                f.write("-" * 40 + "\n")
                recent_logs = log_entries[-20:] if len(log_entries) > 20 else log_entries
                for entry in recent_logs:
                    f.write(f"[{entry['timestamp']}] {entry['level']} - {entry['module']}: {entry['message']}\n")
                
                f.write("\n" + "=" * 80 + "\n")
                f.write("End of Report\n")
                f.write("=" * 80 + "\n")
            
            print(f"📄 TXT report generated: {output_file}")
            return True
        except Exception as e:
            print(f"❌ Error generating TXT report: {e}")
            return False
    
    def generate_word_report(self, output_file, team_info, analysis, log_entries):
        """Generate Word document report"""
        if not DOCX_AVAILABLE:
            print("❌ python-docx not available. Install with: pip install python-docx")
            return False
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('PayBuddy Cybersecurity Assessment Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Team Information
            doc.add_heading('Team Information', level=1)
            team_table = doc.add_table(rows=1, cols=2)
            team_table.style = 'Table Grid'
            
            team_table.cell(0, 0).text = 'Team'
            team_table.cell(0, 1).text = team_info['team']
            
            for member in team_info['members']:
                row = team_table.add_row()
                row.cells[0].text = 'Member'
                row.cells[1].text = f"{member['name']} ({member['reg']})"
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"This report summarizes the cybersecurity assessment conducted on {datetime.now().strftime('%B %d, %Y')}. ")
            summary_para.add_run(f"A total of {analysis['test_summary']['total_tests']} security tests were performed across {len(analysis['modules_used'])} different modules. ")
            
            if analysis['findings']:
                summary_para.add_run(f"{len(analysis['findings'])} security findings were identified. ")
            
            # Test Summary
            doc.add_heading('Test Summary', level=1)
            doc.add_paragraph(f"Total Tests Executed: {analysis['test_summary']['total_tests']}")
            doc.add_paragraph(f"Successful Tests: {analysis['test_summary']['successful_tests']}")
            doc.add_paragraph(f"Modules Used: {', '.join(analysis['modules_used'])}")
            
            # Tests Performed
            if analysis['test_counts']:
                doc.add_heading('Tests Performed', level=2)
                for test_name, count in analysis['test_counts'].items():
                    doc.add_paragraph(f"• {test_name}: {count} execution(s)", style='List Bullet')
            
            # Security Findings
            if analysis['findings']:
                doc.add_heading('Security Findings', level=1)
                for i, finding in enumerate(analysis['findings'], 1):
                    doc.add_heading(f'Finding {i}: {finding["severity"]} Severity', level=2)
                    doc.add_paragraph(f"Timestamp: {finding['timestamp']}")
                    doc.add_paragraph(f"Details: {finding['message']}")
            
            # Errors and Issues
            if analysis['errors']:
                doc.add_heading('Errors and Issues', level=1)
                for error in analysis['errors'][:10]:  # Limit to first 10 errors
                    doc.add_paragraph(f"• [{error['timestamp']}] {error['message']}", style='List Bullet')
            
            # Recommendations
            doc.add_heading('Recommendations', level=1)
            doc.add_paragraph("• Regularly update security testing procedures")
            doc.add_paragraph("• Implement continuous monitoring for identified vulnerabilities")
            doc.add_paragraph("• Conduct regular security awareness training")
            doc.add_paragraph("• Maintain proper logging and incident response procedures")
            
            # Footer
            doc.add_page_break()
            doc.add_paragraph(f"Report generated by PayBuddy Security Toolkit on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Log file integrity hash: {self.logger.calculate_file_hash(self.logger.log_file)[:16]}...")
            
            doc.save(output_file)
            print(f"📄 Word report generated: {output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Error generating Word report: {e}")
            return False
    
    def generate_pdf_report(self, output_file, team_info, analysis, log_entries):
        """Generate PDF report"""
        if not PDF_AVAILABLE:
            print("❌ reportlab not available. Install with: pip install reportlab")
            return False
        
        try:
            doc = SimpleDocTemplate(output_file, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1  # Center
            )
            
            # Title
            story.append(Paragraph("PayBuddy Cybersecurity Assessment Report", title_style))
            story.append(Spacer(1, 20))
            
            # Team Information
            story.append(Paragraph("Team Information", styles['Heading2']))
            team_data = [['Team', team_info['team']]]
            for member in team_info['members']:
                team_data.append(['Member', f"{member['name']} ({member['reg']})"])
            
            team_table = Table(team_data)
            team_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(team_table)
            story.append(Spacer(1, 20))
            
            # Executive Summary
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            summary_text = f"""
            This report summarizes the cybersecurity assessment conducted on {datetime.now().strftime('%B %d, %Y')}.
            A total of {analysis['test_summary']['total_tests']} security tests were performed across 
            {len(analysis['modules_used'])} different modules.
            """
            story.append(Paragraph(summary_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Test Summary
            story.append(Paragraph("Test Summary", styles['Heading2']))
            test_data = [
                ['Metric', 'Value'],
                ['Total Tests', str(analysis['test_summary']['total_tests'])],
                ['Successful Tests', str(analysis['test_summary']['successful_tests'])],
                ['Modules Used', ', '.join(analysis['modules_used'])],
                ['Security Findings', str(len(analysis['findings']))],
                ['Errors Encountered', str(len(analysis['errors']))]
            ]
            
            test_table = Table(test_data)
            test_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            story.append(test_table)
            story.append(Spacer(1, 20))
            
            # Security Findings
            if analysis['findings']:
                story.append(Paragraph("Security Findings", styles['Heading2']))
                for i, finding in enumerate(analysis['findings'], 1):
                    story.append(Paragraph(f"Finding {i}: {finding['severity']} Severity", styles['Heading3']))
                    story.append(Paragraph(f"Timestamp: {finding['timestamp']}", styles['Normal']))
                    story.append(Paragraph(f"Details: {finding['message']}", styles['Normal']))
                    story.append(Spacer(1, 10))
            
            # Footer
            story.append(Spacer(1, 30))
            story.append(Paragraph(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            story.append(Paragraph(f"PayBuddy Security Toolkit v1.0", styles['Normal']))
            
            doc.build(story)
            print(f"📄 PDF report generated: {output_file}")
            return True
            
        except Exception as e:
            print(f"❌ Error generating PDF report: {e}")
            return False

def main():
    parser = argparse.ArgumentParser(description='PayBuddy Logging & Reporting Tool')
    parser.add_argument('--log-event', nargs=3, metavar=('TYPE', 'MODULE', 'MESSAGE'),
                       help='Log an event (type module message)')
    parser.add_argument('--generate-report', choices=['json', 'docx', 'pdf', 'txt'],
                       help='Generate summary report in specified format')
    parser.add_argument('-o', '--output', help='Output file for report')
    parser.add_argument('--view-logs', action='store_true', help='View recent log entries')
    parser.add_argument('--check-integrity', action='store_true', help='Check log file integrity')
    parser.add_argument('--filter', help='Filter log entries by keyword')
    
    args = parser.parse_args()
    
    # Initialize logger
    logger = SecurityLogger()
    
    print("📝 PayBuddy Logging & Reporting Tool")
    print("=" * 50)
    
    # Check integrity
    if args.check_integrity:
        current_hash = logger.calculate_file_hash(logger.log_file)
        try:
            with open(logger.integrity_file, 'r') as f:
                stored_hash = f.read().strip()
            
            if current_hash == stored_hash:
                print("✅ Log file integrity verified")
            else:
                print("❌ Log file integrity check failed!")
                print(f"   Current hash:  {current_hash[:32]}...")
                print(f"   Expected hash: {stored_hash[:32]}...")
        except FileNotFoundError:
            print("⚠️  Integrity file not found")
        
        return 0
    
    # Log an event
    if args.log_event:
        event_type, module, message = args.log_event
        logger.log_event(event_type, message, module)
        print(f"✅ Event logged: {event_type} in {module}")
        return 0
    
    # View logs
    if args.view_logs:
        entries = logger.get_log_entries(filter_by=args.filter)
        
        if not entries:
            print("📭 No log entries found")
            return 0
        
        print(f"📋 Recent Log Entries ({len(entries)} total):")
        print("-" * 80)
        
        # Show last 20 entries
        for entry in entries[-20:]:
            print(f"{entry['timestamp']} | {entry['level']} | {entry['module']} | {entry['message']}")
        
        if len(entries) > 20:
            print(f"... and {len(entries) - 20} more entries")
        
        return 0
    
    # Generate report
    if args.generate_report:
        if not args.output:
            # Auto-generate filename with team info
            verifier = IdentityVerifier()
            team_info = verifier.get_team_info()
            
            if team_info['members']:
                reg_num = team_info['members'][0]['reg'].split('-')[1] if '-' in team_info['members'][0]['reg'] else '000'
                name = team_info['members'][0]['name'].replace(' ', '')
                team_name = team_info['team'].replace(' ', '')
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                args.output = f"report_{team_name}_{reg_num}_{name}_{timestamp}.{args.generate_report}"
            else:
                args.output = f"security_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{args.generate_report}"
        
        # Ensure output has correct extension
        output_path = Path(args.output)
        if output_path.suffix != f'.{args.generate_report}':
            output_path = output_path.with_suffix(f'.{args.generate_report}')
        
        generator = ReportGenerator(logger)
        success = generator.generate_summary_report(output_path, args.generate_report)
        
        if success:
            print(f"✅ Report generation completed")
            return 0
        else:
            print(f"❌ Report generation failed")
            return 1
    
    # Default action - show help
    print("ℹ️  Use --help to see available options")
    print("   Example: python logger.py --view-logs")
    print("   Example: python logger.py --generate-report json -o report.json")
    print("   Example: python logger.py --log-event TEST_START port_scanner 'Starting port scan'")
    
    return 0

if __name__ == "__main__":
    sys.exit(main())